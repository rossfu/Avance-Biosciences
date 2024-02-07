#For quickly generating Biodistribution final reports
#Developed by Ross Fu 4/7/22 -
#ericrossfu@yahoo.com

import os
import sys
import time
import docx
import pwinput
import numpy as np
import pandas as pd
import xlwings as xw
from datetime import datetime
import os.path, mysql.connector

from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


    
    
def run_log(message):
    with open(Final_Report_Output_Path + '\\' + Project_ID + '_' + Runtime + 'AutoReport_run_log.txt', 'a') as z:
        z.write(message)


def setup_stuff():

    #Run Log
    with open(Final_Report_Output_Path + '\\' + Project_ID + '_' + Runtime + 'AutoReport_run_log.txt', 'w') as z:
        z.write('\n#########################################################################\n\n')
        z.write('User: ' + os.environ['USERNAME'] + '\n')
        z.write('Computer Name: ' + os.environ['COMPUTERNAME'] + '\n\n')


    #Validation
    import platform
    import pkg_resources


    if platform.system() != 'Windows':
        print('')
        print('This program is only validated on Windows OS, You are using ' + str(platform.system()))
        run_log('This program is only validated on Windows OS, You are using ' + str(platform.system()))
        print('')

    if platform.python_version() != '3.9.6':
        print('')
        print('This program is only validated on Python ver 3.9.6, You are using ' + str(platform.python_version()))
        run_log('This program is only validated on Python ver 3.9.6, You are using ' + str(platform.python_version()))
        print('')

    if xw.__version__ != '0.25.3':
        print('')
        print('This program is only validated on XLwings ver 0.25.3, You are using ' + str(xw.__version__))
        run_log('This program is only validated on XLwings ver 0.25.3, You are using ' + str(xw.__version__))
        print('')

    if pkg_resources.get_distribution("DateTime").version != '4.3':
        print('')
        print('This program is only validated on DateTime version 4.3, You are using ' + str(pkg_resources.get_distribution("DateTime").version))
        run_log('This program is only validated on DateTime version 4.3, You are using ' + str(pkg_resources.get_distribution("DateTime").version))
        print('')

    if pd.__version__ != '1.3.1':
        print('')
        print('This program is only validated on Pandas ver 1.3.1, You are using pandas ' + str(pd.__version__))
        run_log('This program is only validated on Pandas ver 1.3.1, You are using pandas ' + str(pd.__version__))
        print('')

    if pwinput.__version__ != '1.0.2':
        print('')
        print('This program is only validated on pwinput ver 1.0.2, You are using pwinput ' + str(pwinput.__version__))
        run_log('This program is only validated on pwinput ver 1.0.2, You are using pwinput ' + str(pwinput.__version__))
        print('')

    print('')
    print('##########################################################################\n\n\nSystem validated, Running program')
    run_log('System validated, Running program\n\n' + time.asctime() + '\n\n\n')



    if execute_style == 'VBA':
        print('Active workbook: ' + str(xw.books.active))



##    #Do we have the right workbook?
##    target_wb = xw.books.active
##    if str(target_wb) != '<Book [' + Excel_workbook_that_ran_program + ']>':
##        print('')
##        print('')
##        print('Your active workbook is: ' + str(target_wb))
##        print('This program can only save results to the workbook that ran the program')
##        print('')
##        print('Either make ' + Excel_workbook_that_ran_program + ' your active workbook. (Run this program from ' + Excel_workbook_that_ran_program + ')')
##        run_log('\nEither make ' + Excel_workbook_that_ran_program + ' your active workbook. (Run this program from ' + Excel_workbook_that_ran_program + ')')
##        print('Or change the Excel File Name parameter in Form2')
##        run_log('Or change the Excel File Name parameter in Form2')
##        print('')
##        print('')
##        input('Program will stop. Please try again')
##        os._exit(0)




def authenticateLIMSLogin(): #check login credentials for the database; returns True if login is successful and False if the user quits; also returns last inputted login credentials

    result = False

    print('')
    print('Welcome, please provide your database credentials.')
    run_log('Processing credentials')
    print('')

    while True: #login loop

        username = input("Input username or 'quit': ")

        if (username.lower() == "quit"):
            pword = None
            break

        pword = pwinput.pwinput()
        
        try:
            myDB = mysql.connector.connect(host="192.168.21.69", user=username, password=pword, database="avancelims")

        except:
            print("\nWrong username or password. Try again.\n")
            run_log('\nWrong username or password. Try again.\n')

        else:
            myDB.close()
            result = True
            print("\nCredentials Approved\nFetching Samples from LIMS..\n\n")
            run_log('\nCredentials Approved\nFetching Samples from LIMS..\n\nProject_ID is: ' + Project_ID + '\n\n\n')
            break

    return (result,username,pword)




def get_LIMS_Sample_Information(username, pword): #returns sample list from the database (list of tuples)

    myDB = mysql.connector.connect(host="192.168.21.69", user=username, password=pword, database="avancelims")
    newCursor = myDB.cursor()
    try:
###########################################################################################################################################################################################################################
        query = "SELECT cnumber, animalid, animalgroup, datecollected, samplename, tubelabel FROM samples WHERE idprojects = " + str(Project_ID)
###########################################################################################################################################################################################################################
        newCursor.execute(query)
        result = newCursor.fetchall()
    except:
        print('Actually, the ProjectID you provided resulted in a Error upon querying the database for samples')
        run_log('Actually, the ProjectID you provided resulted in a Error upon querying the database for samples')
        print('')
        print('')
        print('Done..')
        input('')
        os._exit(0)

    newCursor.close()
    myDB.close()

    return result



def get_Appendix1():

    print('Writing Appendix 1')
    run_log('Writing Appendix 1')

    Form1_unordered = pd.read_excel(SCS19_Path, sheet_name = ' Form1_Sample', header=1, usecols='B,F,G,H,C,D', index_col = None, converters={'Animal #': str, 'Group': str})#usecols=['C #', 'Animal #', 'Group', 7, 'Specimen Abbreviation ', 'Tube Label'], index_col=1)#usecols=[1,5,6,7,2,3], index_col = 1)  #usecols=[0,1,2,3,5,6])#, index_col = 0)
    Form1 = Form1_unordered[Form1_unordered.columns[[0,3,4,5,1,2]]]           #2,3,4,0,1]]]
##    Form1 = Form1_unordered

    if execute_style == 'VBA':
        xw.sheets('Report Tables').range('B1').value = 'Appendix 1. Sample Information'
        xw.sheets('Report Tables').range('A2').value = Form1
        xw.sheets('Report Tables').range('A1:A1000').value = ' '

    #Floats into integers
##    for x in range(len(Form1)):
##        if np.isnan(Form1['Animal #'][x]) == False:
##            Form1.iloc[x,1] = int(Form1.iloc[x,1])   #Animal #
##        if np.isnan(Form1['Group'][x]) == False:
##            Form1.iloc[x,2] = int(Form1['Group'][x])         #Group  #



    #Into Report
    Apndx1_title_paragraph = report.add_paragraph()
    Apndx1_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Apndx1_title_run = Apndx1_title_paragraph.add_run('Appendix 1: Tissue Specimens for Biodistribution Testing')
    Apndx1_title_run.bold = True
    report.add_paragraph() #skipline
    report.add_paragraph() #skipline

    
    Apndx1 = report.add_table(Form1.shape[0]+1, Form1.shape[1], style='Final Report Table')
    widths = [Inches(.4),Inches(.5),Inches(.4),Inches(.7),Inches(.7),Inches(3.5)]
##    Apndx1.allow_autofit


    #Header
    for c in range(Form1.shape[-1]):
        Apndx1.cell(0,c).text = Form1.columns[c]
        Apndx1.cell(0,c).width = widths[c]
##        Apndx1.cell(0,c).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    #Bold Header
    for cell in Apndx1.rows[0].cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        run.bold = True

        
    #DF Content
##    for i in range(Form1.shape[0]):
##        for j in range(Form1.shape[-1]):
##            Apndx1.cell(i+1,j).text = str(Form1.values[i,j])

##    values = Form1.values
##    for i in range(Form1.shape[0]):
##        row = Apndx1.add_row()
##        for j, cell in enumerate(row.cells):
##            cell.text = str(values[i, j])

    #Write Values
    Apndx1_cells = Apndx1._cells
    for i in range(Form1.shape[0]):
        for j in range(Form1.shape[1]):
            Apndx1_cells[j + (i+1) * Form1.shape[1]].text = str(Form1.values[i][j])
            Apndx1_cells[j + (i+1) * Form1.shape[1]].width = widths[j]

    #Kill Indents
    for row in Apndx1.rows[:]:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            paragraph.paragraph_format.left_indent = Inches(0)


##    for i in range(Form1.shape[0]):
##        for j, cell in enumerate(Apndx1.rows[i].cells):
##            cell.width = widths[j]
    
    #report.save(Final_Report_Path)
    print('Appendix 1 is in the report')

    

def get_Appendix2():

    print('Writing Appendix 2')
    run_log('Writing Appendix 2')

    Form11 = pd.read_excel(SCS19_Path, sheet_name = 'Form11_Analysis', header=0, usecols=[0,1,2,4,8,9,10,11,12,13], index_col = None, converters={'Animal #': str, 'Group': str, 'gDNA Mass per Reaction (ng)': str})

    if execute_style == 'VBA':
        xw.sheets('Report Tables').range('I1').value = 'Appendix 2. Biodistribution Summary Results'
        xw.sheets('Report Tables').range('H2').value = Form11
        xw.sheets('Report Tables').range('H1:H1000').value = ' '
   
    
    #Into Report
    report.add_page_break()
    Apndx2_title_paragraph = report.add_paragraph()
##    Apndx2_title_paragraph.page_break_before = True
    Apndx2_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Apndx2_title_run = Apndx2_title_paragraph.add_run('Appendix 2: Biodistribution Summary Results')
    Apndx2_title_run.bold = True
    report.add_paragraph() #skipline
    report.add_paragraph() #skipline


    
    Apndx2 = report.add_table(Form11.shape[0]+1, Form11.shape[1], style='Final Report Table')
    widths = [Inches(.6),Inches(.3),Inches(.3),Inches(.5),Inches(1),Inches(1),Inches(.4),Inches(.4),Inches(.5),Inches(1)]
##    Apndx2.allow_autofit

    #Header
    for c in range(Form11.shape[-1]):
        Apndx2.cell(0,c).text = Form11.columns[c]
        Apndx2.cell(0,c).width = widths[c]
        
##        Apndx2.cell(0,c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #Bold Header
    for cell in Apndx2.rows[0].cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        run.bold = True

    
    #DF Content
##    for i in range(Form11.shape[0]):
##        for j in range(Form11.shape[-1]):
##            Apndx2.cell(i+1,j).text = str(Form11.values[i,j])

##    values = Form11.values
##    for i in range(Form11.shape[0]):
##        row = Apndx2.add_row()
##        for j, cell in enumerate(row.cells):
##            cell.text = str(values[i, j])

##    Apndx2.style = 'Final Report Table'

    Apndx2_cells = Apndx2._cells
    for i in range(Form11.shape[0]):
        for j in range(Form11.shape[1]):
            Apndx2_cells[j + (i+1) * Form11.shape[1]].text = str(Form11.values[i][j])
            Apndx2_cells[j + (i+1) * Form11.shape[1]].width = widths[j]

##    for i in range(Form11.shape[0]):
##        for j in range(Form11.shape[1]):
##            Apndx2.cell(i, j).width = widths[j] #10+ min run and doesnt work

    #Kill Indents
    for row in Apndx2.rows[:]:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            paragraph.paragraph_format.left_indent = Inches(0)

    #Apply to every cell func
##    for row in Apndx2.rows[1:]:
##        for cell in row.cells:
##            paragraphs = cell.paragraphs
##            paragraph = paragraphs[0]
##            run_obj = paragraph.runs
##            run = run_obj[0]
##            font = run.font
##            font.size = Pt(3)
    
    #report.save(Final_Report_Path)
    print('Appendix 2 is in the report')



def get_Tissue_above_LOQ_table():

    print('Writing Table 6')
    run_log('Writing Table 6')


    if execute_style == 'VBA':
        xw.sheets('Report Tables').range('S1').value = 'Table 6. Tissues with a Measured Vector Copy Numbers Greater Than LOQ'
##    xw.sheets('Report Tables').range('S2').value = 'Tissue Type'
##    xw.sheets('Report Tables').range('T2').value = 'Number of >LOQ Samples'
##    xw.sheets('Report Tables').range('U2').value = 'Vector Copies Detected'

    Form11 = pd.read_excel(SCS19_Path, sheet_name = 'Form11_Analysis', usecols = [5,9], header=0, index_col = None)


##    Table6 = pd.DataFrame(columns=['Tissue Type', 'Number of >LOQ Samples', 'Vector Copies Detected'])
##    unique_tissue_list = []

    Table6 = {}



    for i in range(len(Form11)):
        if type(Form11['<LOQ '][i]) == float and pd.isna(Form11['Sample Name'][i]) == False: #Its >LOQ
            
            if Form11['Sample Name'][i] not in Table6.keys():
                Table6[Form11['Sample Name'][i]] = 1
            else:
                Table6[Form11['Sample Name'][i]] = Table6.get(Form11['Sample Name'][i]) + 1



##    for i in range(len(Form11)):
##        if type(Form11.iloc[i,1]) == float:# and np.isnan(Form11.iloc[i,0]) == False: #Its >LOQ
##            if Form11.iloc[i,0] not in Table6.keys():
##                Table6[Form11.iloc[i,0]] = 1
##            else:
##                Table6[Form11.iloc[i,0]] = Table6.get(Form11.iloc[i,0]) + 1
           

    dictionary = [Table6]
    Table6_df = pd.DataFrame()
    Table6_df['Tissue Type'] = Table6.keys()
    Table6_df['Number of >LOQ Samples'] = Table6.values()
    Table6_df['Vector Copies Detected'] = '>LOQ'
    Table6_df.reset_index(drop=True, inplace=True)

    if execute_style == 'VBA':
        xw.sheets('Report Tables').range('R2').value = Table6_df
        xw.sheets('Report Tables').range('R1:R100').value = ' '
        
            
##            if Form11['Sample Name'][row] not in Table6:
##                Table6[Form11['Sample Name'][row]] = 1
##            else:
##                Table6[Form11['Sample Name'][row]] =


    #Into Report
    Table6 = report.add_table(Table6_df.shape[0]+1, Table6_df.shape[1], style='Table Grid')

    #Header
    for c in range(Table6_df.shape[-1]):
        Table6.cell(0,c).text = Table6_df.columns[c]

    #Bold Header
    for cell in Table6.rows[0].cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        run.bold = True
        
    #DF Content
##    for i in range(Table6_df.shape[0]):
##        for j in range(Table6_df.shape[-1]):
##            Table6.cell(i+1,j).text = str(Table6_df.values[i,j])

##    values = Table6_df.values
##    for i in range(Table6_df.shape[0]):
##        row = Table6.add_row()
##        for j, cell in enumerate(row.cells):
##            cell.text = str(values[i, j])

    Table6_cells = Table6._cells
    for i in range(Table6_df.shape[0]):
        for j in range(Table6_df.shape[1]): #Offset 1 for header
            Table6_cells[j + (i+1) * Table6_df.shape[1]].text = str(Table6_df.values[i][j])

    Table6_location = False

    for x in report.paragraphs:
        if 'particles over time.' in x.text:
            Table6_location = True
            Table6_preceding_paragraph = x._p

            Table6_title_paragraph = report.add_paragraph()
            Table6_preceding_paragraph.addnext(Table6_title_paragraph._p)
            Table6_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Table6_title_run = Table6_title_paragraph.add_run('Table 6: Tissues Above LOQ')
            Table6_title_run.bold = True

            
            z = Table6_title_paragraph._p
            z.addnext(Table6._tbl)

    if Table6_location == False:
        print('Could not locate Table 6 Location, not added')
        run_log('Could not locate Table 6 Location, not added')
    
    
    report.save(Final_Report_Path)
    os.startfile(Final_Report_Path)
    
##    tbl, p = table6._tbl, Table6_preceding_paragraph._p
##    p.addnext(tbl)
    
    print('Table 6 is in the report')
    run_log('Table 6 is in the report')



def func_write_information_to_Excel_workbook(title, content, which_column_letters): #Only used for Form1_LIMS

    currentRow = 3
    currentCol = 0

    if execute_style == 'VBA':
        if title.startswith('Appendix 1'):

            print('Writing Appendix 1')
            run_log('Writing Appendix 1')

            xw.sheets('Report Tables').range('A1').value = title
            xw.sheets('Report Tables').range('A2').value = 'Sample ID'
            xw.sheets('Report Tables').range('B2').value = 'Animal #'
            xw.sheets('Report Tables').range('C2').value = 'Group'
            xw.sheets('Report Tables').range('D2').value = 'Collection Time Point'
            xw.sheets('Report Tables').range('E2').value = 'Sample Type'
            xw.sheets('Report Tables').range('F2').value = 'Tube Label'


        
        for row in content:
            for fieldIndex in range(len(which_column_letters)):
                

                currentCol_letter = which_column_letters[currentCol]
                currentCell = xw.sheets('Report Tables').range(str(currentCol_letter)+str(currentRow))

                if (row[fieldIndex] and row[fieldIndex] != "\n"): #only try to copy field if it has a value
                    if (row[fieldIndex][0].isdigit() or row[fieldIndex].startswith("-")): #set cell value as a float if it begins with a number or a "-"

                        try: #copy value as a float
                            currentCell.value = float(row[fieldIndex])

                        except: #copy value as text
                            currentCell.value = row[fieldIndex].strip()

                    else: #copy value as text
                        currentCell.value = row[fieldIndex].strip()

                if currentCol < len(which_column_letters)-1:
                    currentCol += 1
                else:
                    currentCol = 0
            currentRow += 1


    

    
def execute():

    Form1='SCS19'


    #Set Up
    setup_stuff()


    #Authenticate
    if Form1 == 'LIMS':    
        credentials = authenticateLIMSLogin()
        loginSuccess = credentials[0]
        if loginSuccess:
##            credentials = ['fill', 'rfu', 'M0nkey!!']
    #Tables
            Sample_Information = get_LIMS_Sample_Information(credentials[1], credentials[2])
            func_write_information_to_Excel_workbook('Appendix 1. Sample Information', Sample_Information, ['A','B','C','D','E','F'])


    elif Form1 == 'SCS19':
        Sample_Information = get_Appendix1()
                  
    get_Appendix2()

    get_Tissue_above_LOQ_table()

    print('Done!')
    
    



def main(a,b,c,d,e,f,g):

    global Project_ID
    global SCS19_Path
    global Final_Report_Output_Path
    global LIMS_UserID
    global LIMS_pw
    global Final_Report_Template_Path
    global Excel_workbook_that_ran_program
    global Runtime

    global report
    
    Project_ID = a
    SCS19_Path = b
    Final_Report_Output_Path = c
    LIMS_UserID = d
    LIMS_pw = e
    Final_Report_Template_Path = f
    Excel_workbook_that_ran_program = g

    now = datetime.now()
    Runtime = now.strftime('%m_%d_%Y__time_%H_%M')

    report_template = docx.Document(Final_Report_Template_Path)

    global ymdhm
    ymdhm = now.strftime('%Y_%m_%d_%H_%M')
    
    global Final_Report_Path
    Final_Report_Path = Final_Report_Output_Path + '\\' + 'Docx_generated_final_report_' + str(ymdhm) + '.docx'

    report_template.save(Final_Report_Path)

    report = docx.Document(Final_Report_Path)



    
    try:
        execute()

    except BaseException as e: #Error Handling

        print('')
        print('######################################################################################')
        print('')
        print('Runtime Exception Occured, saving into error log file')
        print('')
        print('')
        
        import traceback

        exc_type, exc_value, exc_traceback = sys.exc_info()

        print(''.join(traceback.format_exception(exc_type, exc_value, exc_traceback)))

        print('#######################################################################################')
        print('')
        print('')
        print('Results were not saved into a Final Report because the run did not complete \n\n\n\n\n###########################\n\n\n')

        run_log('\n#######################################################################################\n\n' + time.asctime() + '\n\n')
        run_log(''.join(traceback.format_exception(exc_type, exc_value, exc_traceback)))
        run_log('\n\n\n\n#######################################################################################\n\n\n\n')
        run_log('Results were not saved into a Final Report because the run did not complete')




#####################################################################
        
execute_style = 'VBA'

if execute_style == 'Python':
    a = '214144'                                                                                                                    ##    Project_ID
    b = r'C:\Users\efu\Desktop\Python\(v2.0) CNG214057_ SCS 19.01 Biodistribution Study unlocked.xlsx'                              ##    SCS19_Path
    c = r'C:\Users\efu\Desktop\Report Test'                                                                                         ##    Final_Report_Output_Path
    d = 'rfu'                                                                                                                       ##    LIMS_UserID
    e = 'M0nkey!!'                                                                                                                  ##    LIMS_pw
    f = r'\\Birch\Operations\5 GlobalShare\Python Scripts\AutoReport\Sample Reports\Biodistribution_Final_Report_Template.docx'     ##    Final_Report_Template_Path
    g = r'\\Birch\Operations\5 GlobalShare\Python Scripts\AutoReport\AutoReport.xlsm'                                               ##    Excel_workbook_that_ran_program

    main(a,b,c,d,e,f,g)
